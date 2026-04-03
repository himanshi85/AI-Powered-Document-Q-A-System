import json
import os
import uuid
from pathlib import Path
from typing import List, Tuple

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader

from .chunker import SemanticChunker
from .models import Document, DocumentChunk


class DocumentProcessor:
    def __init__(self, target_tokens: int = 350, overlap_tokens: int = 60):
        self.chunker = SemanticChunker(
            target_tokens=target_tokens,
            overlap_tokens=overlap_tokens,
        )
        self.last_processing_report = []

    def _detect_doc_type(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        mapping = {
            ".txt": "text",
            ".md": "markdown",
            ".json": "json",
            ".csv": "csv",
            ".pdf": "pdf",
            ".docx": "docx",
        }
        return mapping.get(ext, "unknown")

    def _clean_text(self, text: str) -> str:
        if text is None:
            return ""
        text = str(text).replace("\x00", " ")
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        return text.strip()

    def _read_txt_md(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return self._clean_text(f.read())

    def _read_json(self, file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            data = json.load(f)
        return self._clean_text(json.dumps(data, indent=2, ensure_ascii=False))

    def _read_csv(self, file_path: str) -> str:
        df = pd.read_csv(file_path)
        return self._clean_text(df.to_csv(index=False))

    def _read_pdf(self, file_path: str) -> Tuple[str, dict]:
        """
        Read PDF text page by page.
        Returns:
            text, metadata
        """
        reader = PdfReader(file_path)
        pages = []
        extracted_pages = 0
        total_chars = 0

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""

            page_text = self._clean_text(page_text)

            if page_text:
                extracted_pages += 1
                total_chars += len(page_text)
                pages.append(f"[Page {i + 1}]\n{page_text}")

        metadata = {
            "pdf_total_pages": len(reader.pages),
            "pdf_extracted_pages": extracted_pages,
            "pdf_total_extracted_chars": total_chars,
        }

        return "\n\n".join(pages), metadata

    def _read_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return self._clean_text("\n\n".join(paragraphs))

    def _read_file(self, file_path: str) -> Tuple[str, dict]:
        ext = Path(file_path).suffix.lower()

        if ext in [".txt", ".md"]:
            return self._read_txt_md(file_path), {}

        if ext == ".json":
            return self._read_json(file_path), {}

        if ext == ".csv":
            return self._read_csv(file_path), {}

        if ext == ".pdf":
            return self._read_pdf(file_path)

        if ext == ".docx":
            return self._read_docx(file_path), {}

        raise ValueError(f"Unsupported file type: {ext}")

    def load_single_document(self, file_path: str) -> Document:
        file_path = str(file_path)
        text, extra_metadata = self._read_file(file_path)

        title = Path(file_path).stem
        doc_type = self._detect_doc_type(file_path)

        return Document(
            document_id=str(uuid.uuid4()),
            title=title,
            text=text,
            source_path=file_path,
            doc_type=doc_type,
            metadata={
                "file_name": Path(file_path).name,
                "file_extension": Path(file_path).suffix.lower(),
                "title": title,
                **extra_metadata,
            },
        )

    def load_documents_from_directory(self, directory_path: str) -> List[Document]:
        supported_extensions = {".txt", ".md", ".json", ".csv", ".pdf", ".docx"}
        documents: List[Document] = []
        self.last_processing_report = []

        for root, _, files in os.walk(directory_path):
            for file_name in files:
                file_path = os.path.join(root, file_name)

                if Path(file_path).suffix.lower() not in supported_extensions:
                    self.last_processing_report.append(
                        {
                            "file": file_name,
                            "status": "skipped",
                            "reason": "unsupported_extension",
                        }
                    )
                    continue

                try:
                    document = self.load_single_document(file_path)

                    if document.text.strip():
                        documents.append(document)
                        self.last_processing_report.append(
                            {
                                "file": file_name,
                                "status": "loaded",
                                "chars": len(document.text),
                                "doc_type": document.doc_type,
                                "metadata": document.metadata,
                            }
                        )
                    else:
                        self.last_processing_report.append(
                            {
                                "file": file_name,
                                "status": "skipped",
                                "reason": "empty_extracted_text",
                                "doc_type": document.doc_type,
                                "metadata": document.metadata,
                            }
                        )

                except Exception as e:
                    self.last_processing_report.append(
                        {
                            "file": file_name,
                            "status": "error",
                            "reason": str(e),
                        }
                    )

        return documents

    def process_documents(self, documents: List[Document]) -> List[DocumentChunk]:
        all_chunks: List[DocumentChunk] = []

        for document in documents:
            chunks = self.chunker.chunk_document(document)
            all_chunks.extend(chunks)

        return all_chunks

    def process_directory(self, directory_path: str) -> List[DocumentChunk]:
        documents = self.load_documents_from_directory(directory_path)
        return self.process_documents(documents)